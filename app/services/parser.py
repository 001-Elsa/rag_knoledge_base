"""Structure-aware parsing and cleaning for uploaded RAG sources.

``parse_file`` keeps the original ``[(text, page)]`` contract.  New ingestion code
uses ``parse_document`` to retain sections, table/image markers, and source URLs.
Heavy/optional parsers are imported lazily so API startup remains cheap.
"""

import csv
import io
import re
import unicodedata
from collections import Counter
from dataclasses import dataclass
from pathlib import Path
from zipfile import BadZipFile, ZipFile

from app.config import settings

SUPPORTED_EXTENSIONS = {
    ".pdf",
    ".docx",
    ".xlsx",
    ".xlsm",
    ".csv",
    ".md",
    ".txt",
    ".html",
    ".htm",
    ".png",
    ".jpg",
    ".jpeg",
    ".webp",
    ".tif",
    ".tiff",
}
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".webp", ".tif", ".tiff"}

Segment = tuple[str, int | None]


@dataclass(frozen=True)
class ParsedSegment:
    text: str
    page: int | None = None
    section: str | None = None
    content_type: str = "text"
    source_url: str | None = None


_CONTROL_CHARS = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]")
_BOILERPLATE = re.compile(
    r"^(?:copyright|all rights reserved|版权所有|扫码关注|点击下载|返回顶部|"
    r"cookie policy|privacy policy)\b",
    re.IGNORECASE,
)


def clean_text(text: str, *, deduplicate: bool = True) -> str:
    """Normalize Unicode/whitespace, remove controls, boilerplate and duplicates."""
    cjk_punctuation = "，。！？；：、“”‘’（）【】《》"
    protected = {character: chr(0xE000 + index) for index, character in enumerate(cjk_punctuation)}
    text = text or ""
    for character, marker in protected.items():
        text = text.replace(character, marker)
    text = unicodedata.normalize("NFKC", text)
    for character, marker in protected.items():
        text = text.replace(marker, character)
    text = text.replace("\ufeff", "").replace("\ufffd", "")
    text = _CONTROL_CHARS.sub("", text)
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    lines = []
    seen: set[str] = set()
    for raw in text.splitlines():
        line = raw.strip()
        if not line or _BOILERPLATE.match(line):
            if lines and lines[-1] != "":
                lines.append("")
            continue
        key = re.sub(r"\s+", "", line).casefold()
        if deduplicate and len(key) >= 12 and key in seen:
            continue
        if len(key) >= 12:
            seen.add(key)
        lines.append(line)
    return re.sub(r"\n{3,}", "\n\n", "\n".join(lines)).strip()


def parse_file(filepath: str) -> list[Segment]:
    """Backward-compatible parser result used by older callers and tests."""
    return [(item.text, item.page) for item in parse_document(filepath)]


def parse_document(filepath: str, *, source_url: str | None = None) -> list[ParsedSegment]:
    path = Path(filepath)
    ext = path.suffix.lower()
    if ext == ".pdf":
        segments = _parse_pdf(path)
    elif ext == ".docx":
        segments = _parse_docx(path)
    elif ext in {".xlsx", ".xlsm"}:
        segments = _parse_excel(path)
    elif ext == ".csv":
        segments = _parse_csv(path)
    elif ext == ".md":
        segments = _parse_markdown(path.read_text(encoding="utf-8", errors="replace"))
    elif ext == ".txt":
        segments = [ParsedSegment(clean_text(path.read_text(encoding="utf-8", errors="replace")))]
    elif ext in {".html", ".htm"}:
        segments = _parse_html(path.read_text(encoding="utf-8", errors="replace"))
    elif ext in IMAGE_EXTENSIONS:
        segments = [_parse_image(path)]
    else:
        raise ValueError(f"不支持的文件类型: {ext}（支持 {sorted(SUPPORTED_EXTENSIONS)}）")
    if source_url:
        segments = [
            ParsedSegment(s.text, s.page, s.section, s.content_type, source_url)
            for s in segments
        ]
    return [s for s in segments if s.text.strip()]


def _parse_pdf(path: Path) -> list[ParsedSegment]:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    if len(reader.pages) > settings.max_document_pages:
        raise ValueError(
            f"PDF 页数超过限制（{len(reader.pages)} > {settings.max_document_pages}）"
        )
    page_texts = [page.extract_text() or "" for page in reader.pages]
    page_texts = _remove_repeated_margins(page_texts)
    results: list[ParsedSegment] = []
    for page_number, text in enumerate(page_texts, start=1):
        cleaned = clean_text(text)
        if len(cleaned) < settings.ocr_min_text_chars and settings.ocr_enabled:
            ocr_text = _ocr_pdf_page(path, page_number - 1)
            if len(ocr_text) > len(cleaned):
                cleaned = ocr_text
        tables = _pdf_tables(path, page_number - 1)
        if tables:
            cleaned = f"{cleaned}\n\n{tables}".strip()
        results.append(ParsedSegment(cleaned, page_number, content_type="page"))
    return results


def _remove_repeated_margins(pages: list[str]) -> list[str]:
    """Remove header/footer lines repeated on most pages (requires >= 3 pages)."""
    if len(pages) < 3:
        return pages
    candidates: Counter[str] = Counter()
    page_lines: list[list[str]] = []
    for page in pages:
        lines = [line.strip() for line in page.splitlines() if line.strip()]
        page_lines.append(lines)
        for line in set(lines[:1] + lines[-1:]):
            key = re.sub(r"\d+", "#", re.sub(r"\s+", " ", line)).casefold()
            if 2 <= len(key) <= 120:
                candidates[key] += 1
    repeated = {key for key, count in candidates.items() if count / len(pages) >= 0.6}
    cleaned_pages = []
    for lines in page_lines:
        kept = []
        for line in lines:
            key = re.sub(r"\d+", "#", re.sub(r"\s+", " ", line)).casefold()
            if key not in repeated:
                kept.append(line)
        cleaned_pages.append("\n".join(kept))
    return cleaned_pages


def _ocr_image(image) -> str:
    try:
        import pytesseract
    except ImportError as exc:
        raise ValueError("OCR 依赖未安装：请安装 pytesseract 与 Tesseract") from exc
    return clean_text(pytesseract.image_to_string(image, lang=settings.ocr_languages))


def _ocr_pdf_page(path: Path, page_index: int) -> str:
    try:
        import fitz
    except ImportError as exc:
        raise ValueError("扫描 PDF OCR 需要 PyMuPDF") from exc
    with fitz.open(path) as document:
        page = document.load_page(page_index)
        pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
        from PIL import Image

        image = Image.frombytes("RGB", (pixmap.width, pixmap.height), pixmap.samples)
        return _ocr_image(image)


def _pdf_tables(path: Path, page_index: int) -> str:
    if not settings.table_extraction_enabled:
        return ""
    try:
        import pdfplumber
    except ImportError:
        return ""
    rendered = []
    with pdfplumber.open(path) as pdf:
        for table in pdf.pages[page_index].extract_tables() or []:
            rows = [[clean_text(cell or "", deduplicate=False) for cell in row] for row in table]
            if rows and any(any(cell for cell in row) for row in rows):
                rendered.append(_markdown_table(rows))
    return "\n\n".join(rendered)


def _parse_image(path: Path) -> ParsedSegment:
    if not settings.ocr_enabled:
        raise ValueError("图片/扫描件解析需要启用 OCR_ENABLED")
    from PIL import Image

    with Image.open(path) as image:
        text = _ocr_image(image.convert("RGB"))
    if not text:
        raise ValueError("OCR 未识别到文字")
    return ParsedSegment(text, page=1, section="图片 OCR", content_type="image_ocr")


def _parse_docx(path: Path) -> list[ParsedSegment]:
    import docx

    try:
        with ZipFile(path) as archive:
            expanded_bytes = sum(info.file_size for info in archive.infolist())
    except BadZipFile as exc:
        raise ValueError("DOCX 压缩包损坏") from exc
    if expanded_bytes > settings.max_uncompressed_mb * 1024 * 1024:
        raise ValueError(f"DOCX 解压大小超过限制（{settings.max_uncompressed_mb}MB）")
    document = docx.Document(str(path))
    segments: list[ParsedSegment] = []
    section = "正文"
    buffer: list[str] = []

    def flush():
        if buffer:
            text = clean_text("\n\n".join(buffer))
            if text:
                segments.append(ParsedSegment(text, section=section))
            buffer.clear()

    for paragraph in document.paragraphs:
        text = clean_text(paragraph.text, deduplicate=False)
        if not text:
            continue
        style = (paragraph.style.name or "").lower() if paragraph.style else ""
        if style.startswith("heading") or style.startswith("标题"):
            flush()
            section = text
        else:
            buffer.append(text)
    flush()
    for index, table in enumerate(document.tables, start=1):
        rows = [[clean_text(cell.text, deduplicate=False) for cell in row.cells] for row in table.rows]
        rendered = _markdown_table(rows)
        if rendered:
            segments.append(ParsedSegment(rendered, section=f"表格 {index}", content_type="table"))
    if settings.ocr_enabled:
        from PIL import Image

        with ZipFile(path) as archive:
            media = [name for name in archive.namelist() if name.startswith("word/media/")]
            for index, name in enumerate(media, start=1):
                try:
                    with Image.open(io.BytesIO(archive.read(name))) as image:
                        text = _ocr_image(image.convert("RGB"))
                    if text:
                        segments.append(
                            ParsedSegment(
                                text,
                                section=f"内嵌图片 {index}",
                                content_type="image_ocr",
                            )
                        )
                except Exception:
                    continue
    return segments


def _parse_excel(path: Path) -> list[ParsedSegment]:
    from openpyxl import load_workbook

    workbook = load_workbook(path, read_only=True, data_only=True)
    results = []
    try:
        for sheet in workbook.worksheets:
            rows = []
            for row in sheet.iter_rows(values_only=True):
                values = [clean_text(str(value), deduplicate=False) if value is not None else "" for value in row]
                if any(values):
                    rows.append(values)
                if len(rows) >= settings.max_spreadsheet_rows:
                    break
            text = _markdown_table(rows)
            if text:
                results.append(ParsedSegment(text, section=sheet.title, content_type="table"))
    finally:
        workbook.close()
    return results


def _parse_csv(path: Path) -> list[ParsedSegment]:
    rows = []
    with path.open("r", encoding="utf-8-sig", errors="replace", newline="") as stream:
        for index, row in enumerate(csv.reader(stream)):
            if index >= settings.max_spreadsheet_rows:
                break
            rows.append([clean_text(cell, deduplicate=False) for cell in row])
    return [ParsedSegment(_markdown_table(rows), section=path.stem, content_type="table")]


def _markdown_table(rows: list[list[str]]) -> str:
    if not rows:
        return ""
    width = max(len(row) for row in rows)
    normalized = [row + [""] * (width - len(row)) for row in rows]
    header = normalized[0]
    body = normalized[1:]
    lines = [" | ".join(header), " | ".join(["---"] * width)]
    lines.extend(" | ".join(row) for row in body)
    return clean_text("\n".join(lines), deduplicate=False)


def _parse_markdown(text: str) -> list[ParsedSegment]:
    results = []
    section = "正文"
    buffer = []
    for line in text.splitlines():
        heading = re.match(r"^#{1,6}\s+(.+)$", line.strip())
        if heading:
            if buffer:
                results.append(ParsedSegment(clean_text("\n".join(buffer)), section=section))
                buffer.clear()
            section = clean_text(heading.group(1), deduplicate=False)
        else:
            buffer.append(line)
    if buffer:
        results.append(ParsedSegment(clean_text("\n".join(buffer)), section=section))
    return results


def _parse_html(text: str) -> list[ParsedSegment]:
    from bs4 import BeautifulSoup

    soup = BeautifulSoup(text, "html.parser")
    for tag in soup(["script", "style", "noscript", "nav", "footer", "aside", "form"]):
        tag.decompose()
    results = []
    section = soup.title.get_text(" ", strip=True) if soup.title else "网页正文"
    buffer = []
    root = soup.find("main") or soup.find("article") or soup.body or soup
    for node in root.find_all(["h1", "h2", "h3", "h4", "p", "li", "table"]):
        if node.name.startswith("h"):
            if buffer:
                results.append(ParsedSegment(clean_text("\n\n".join(buffer)), section=section))
                buffer.clear()
            section = clean_text(node.get_text(" ", strip=True), deduplicate=False)
        else:
            value = clean_text(node.get_text(" | " if node.name == "table" else " ", strip=True), deduplicate=False)
            if value:
                buffer.append(value)
    if buffer:
        results.append(ParsedSegment(clean_text("\n\n".join(buffer)), section=section))
    return results
